"""
文本特征提取器（PDF/DOCX/Markdown/纯文本 → 抽文本 → MinHash 256-bit）

支持 4 种输入：
- application/pdf：pypdf 抽页面文本
- application/vnd.openxmlformats-officedocument.wordprocessingml.document（docx）：python-docx
- text/markdown：markdown 渲染后取纯文本
- text/plain：直接读

抽到的文本经 jieba 分词 + 中英 5-gram shingle，喂给 MinHash 压缩为 256-bit。
SimHash 在短文本和高频词面前失效，MinHash 对集合相似度（Jaccard）有数学保证，
对增删改、句子顺序变化也鲁棒。
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

import jieba
import markdown as md_lib
from datasketch import MinHash
from docx import Document
from pypdf import PdfReader

from app.extractors.base import HASH_BITS, Extractor, ExtractResult, bits_to_hex, extra_json
from app.extractors.download import download_to_tempfile

logger = logging.getLogger(__name__)

SHINGLE_SIZE = 5
MIN_TEXT_CHARS = 10  # 太短文本无意义指纹


# ========== 文本抽取 ==========

def _extract_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:
            logger.warning("PDF 页面抽文本失败: %s", e)
    return "\n".join(parts)


def _extract_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_markdown(path: str) -> str:
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")
    html = md_lib.markdown(raw)
    # 简易剥 HTML 标签
    return re.sub(r"<[^>]+>", " ", html)


def _extract_plain(path: str) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def _read_text(file_path: str) -> str:
    """按文件头/扩展名嗅探类型抽取文本"""
    head = Path(file_path).read_bytes()[:8]
    if head.startswith(b"%PDF"):
        return _extract_pdf(file_path)
    if head.startswith(b"PK\x03\x04"):
        # DOCX 是 ZIP，尝试以 docx 解析
        try:
            return _extract_docx(file_path)
        except Exception as e:
            logger.warning("DOCX 解析失败，回退纯文本: %s", e)
            return _extract_plain(file_path)
    # 其余按文本：md/txt 都走纯文本路径，jieba 分词后效果一致
    text = _extract_plain(file_path)
    if text.lstrip().startswith(("#", "*", "-", "=", "[", ">")):
        # 看起来像 markdown，渲染一次去掉语法符号
        return re.sub(r"<[^>]+>", " ", md_lib.markdown(text))
    return text


# ========== Shingle + MinHash ==========

def _normalize(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _shingles(text: str) -> list[str]:
    """
    生成 5-gram shingle 集合
    - 中文：jieba 分词后按 token 5-gram
    - 英文/混合：分词结果天然包含 ASCII 词，复用同一通道
    """
    tokens = [t for t in jieba.cut(text) if t.strip()]
    if len(tokens) < SHINGLE_SIZE:
        # 短文本：退化为字符 5-gram，仍能产生稳定指纹
        compact = re.sub(r"\s+", "", text)
        return [compact[i:i + SHINGLE_SIZE] for i in range(max(1, len(compact) - SHINGLE_SIZE + 1))]
    return [" ".join(tokens[i:i + SHINGLE_SIZE]) for i in range(len(tokens) - SHINGLE_SIZE + 1)]


class MinhashTextExtractor(Extractor):
    """文本 MinHash 提取器"""

    @property
    def algo(self) -> str:
        return "MINHASH"

    @property
    def algo_version(self) -> str:
        return "1.0"

    def extract(self, file_path: str) -> ExtractResult:
        # 1. 嗅探格式抽出纯文本
        raw = _read_text(file_path)
        text = _normalize(raw)
        if len(text) < MIN_TEXT_CHARS:
            raise RuntimeError(f"文本内容过短（< {MIN_TEXT_CHARS} 字符），无法生成有效指纹")

        # 2. 5-gram shingle 集合
        shingles = _shingles(text)
        if not shingles:
            raise RuntimeError("文本分词后无可用 shingle")

        # 3. MinHash 压缩为 256-bit
        mh = MinHash(num_perm=HASH_BITS, seed=42)
        for s in shingles:
            mh.update(s.encode("utf-8"))
        bits = (mh.hashvalues & 1).astype("uint8")

        return ExtractResult(
            algo=self.algo,
            algo_version=self.algo_version,
            perceptual_hash=bits_to_hex(bits),
            extra={
                "char_count": len(text),
                "shingle_count": len(shingles),
                "shingle_size": SHINGLE_SIZE,
            },
        )


def extract_text(file_url: str) -> dict:
    """主流程入口：下载 → 抽文本 → shingle → MinHash → 返回响应字典"""
    logger.info("开始文本特征提取: %s", file_url[:100])
    extractor = MinhashTextExtractor()
    with download_to_tempfile(file_url, suffix=".doc") as path:
        result = extractor.extract(path)
    logger.info("文本特征提取完成: hash=%s, chars=%s",
                result.perceptual_hash[:16], result.extra.get("char_count"))
    return {
        "work_type": "TEXT",
        "algo": result.algo,
        "algo_version": result.algo_version,
        "perceptual_hash": result.perceptual_hash,
        "extra": extra_json(result.extra),
    }
